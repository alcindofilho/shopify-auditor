import streamlit as st
import requests
from bs4 import BeautifulSoup
import google.generativeai as genai
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import datetime

# --- 1. AGENCY CONFIGURATION ---
AGENCY_NAME = "Inkroast"
AGENCY_URL = "https://www.inkroast.com"
BOOKING_LINK = "https://portal.inkroast.com/discovery"

# --- AFFILIATE LINKS DATABASE ---
# Map the exact app name (as the AI sees it) to your specific affiliate URL
APP_LINKS = {
    "Plug In SEO": "https://pluginseo.com?ref=YOUR_REF_ID",
    "Klaviyo": "https://klaviyo.com/partner/YOUR_REF_ID",
    "Judge.me": "https://judge.me/ref/YOUR_REF_ID",
    "Loox": "https://loox.io/app/YOUR_REF_ID",
    "Privy": "https://privy.com/ref/YOUR_REF_ID",
    "ReConvert": "https://reconvert.io/?ref=YOUR_REF_ID"
}

# --- 2. CONFIGURATION & STYLING ---
st.set_page_config(page_title=f"{AGENCY_NAME} Shopify e-commerce Snapshot", page_icon="☕", layout="centered")

st.markdown("""
<style>
    /* Report Container */
    .report-container {
        background-color: #000000;
        padding: 40px;
        border: 1px solid #e0e0e0;
        border-radius: 8px;
        font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
        color: #333;
    }
    
    /* Typography */
    h1 { color: #1a1a1a; font-family: 'Arial', sans-serif; font-weight: 800; }
    h2 { color: #2c3e50; font-family: 'Arial', sans-serif; font-weight: 600; }

    /* --- INPUT FIELD SHADOW --- */
    /* This targets the actual input box inside Streamlit's wrapper */
    div[data-baseweb="input"] > div {
        box-shadow: 0 4px 10px rgba(0, 0, 0, 0.08) !important;
        border: 1px solid #e0e0e0 !important;
        border-radius: 6px !important;
    }
    
    /* "Generate Report" Button - INKROAST GREEN */
    .stButton>button { 
        width: 100%; 
        background-color: #008060; 
        color: white; 
        font-weight: bold; 
        border: none;
        padding: 14px;
        border-radius: 6px;
        font-size: 16px;
        text-transform: uppercase;
        letter-spacing: 0.5px;
        transition: all 0.3s ease;
    }
    .stButton>button:hover {
        background-color: #004c3f; 
        box-shadow: 0 4px 12px rgba(0,0,0,0.15);
    }
    
    /* Success Message Box */
    .stSuccess {
        background-color: #f0fdf4;
        border: 1px solid #bbf7d0;
        color: #166534;
    }
    
    /* --- TAB STYLING --- */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background-color: transparent;
    }
    
    .stTabs [data-baseweb="tab"] {
        height: 50px;
        white-space: pre-wrap;
        background-color: #f0f2f6; /* GREY Background */
        color: #000000;            /* BLACK Text */
        border-radius: 4px;
        padding: 10px 20px;
        font-weight: 600;
        border: 1px solid #e0e0e0;
    }
    
    /* Active Tab Styling */
    .stTabs [aria-selected="true"] {
        background-color: #e6fffa; 
        color: #000000;            
        border: 1px solid #008060; 
        border-bottom: 3px solid #008060; 
    }
    
    /* Service/Tool Box Styling */
    .service-box {
        background-color: #f0f2f6; 
        padding: 15px; 
        border-radius: 6px; 
        border-left: 5px solid #008060; 
        color: #000000;
        margin-bottom: 20px;
    }
</style>
""", unsafe_allow_html=True)

# --- 3. AUTHENTICATION ---
try:
    api_key = st.secrets["GEMINI_API_KEY"]
    genai.configure(api_key=api_key)
except Exception:
    st.error("⚠️ API Key missing. Please set GEMINI_API_KEY in Streamlit Secrets.")
    st.stop()

# --- 4. CORE FUNCTIONS ---

def scrape_shopify_store(url):
    """Scrapes the homepage for text content AND technical SEO data."""
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'}
        url = url.strip()
        if not url.startswith('http'):
            url = 'https://' + url
            
        response = requests.get(url, headers=headers, timeout=10)
        if response.status_code != 200:
            return None, f"Could not load site (Status: {response.status_code})"
            
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # 1. Basic Content
        title = soup.title.string if soup.title else "No Title"
        meta = soup.find('meta', attrs={'name': 'description'})
        desc = meta['content'] if meta else "No Meta Description"
        
        # 2. Technical SEO Data (H-Tags)
        h1_tags = [h.get_text().strip() for h in soup.find_all('h1')]
        h2_tags = [h.get_text().strip() for h in soup.find_all('h2')]
        
        # 3. Image Analysis (Alt Tags)
        images = soup.find_all('img')
        total_images = len(images)
        missing_alt = sum(1 for img in images if not img.get('alt') or img.get('alt') == "")
        
        body = soup.get_text(separator=' ', strip=True)[:5000]
        
        return {
            "url": url,
            "title": title,
            "description": desc,
            "h1_tags": h1_tags,
            "h2_tags": h2_tags[:6],
            "image_stats": f"{missing_alt} out of {total_images} images are missing description tags (Alt Text).",
            "body": body
        }, None
    except Exception as e:
        return None, f"Scraping Error: {str(e)}"

def analyze_store_json(data):
    """Generates the Inkroast Strategic Audit based on the 6-point framework."""
    model = genai.GenerativeModel('models/gemini-2.5-flash')
    
    prompt = f"""
    You are a Senior Strategist at {AGENCY_NAME} specialized in the coffee industry.
    
    Analyze this store data:
    - URL: {data['url']}
    - Title Tag: {data['title']}
    - Meta Desc: {data['description']}
    - H1 Tags found: {data['h1_tags']}
    - H2 Tags found: {data['h2_tags']}
    - Image Analysis: {data['image_stats']}
    - Content Snippet: {data['body'][:3500]}

    **Your Role:**
    Perform a deep audit covering exactly these 6 sections:
    
    1. **Branding & Messaging:** Clarity, consistency, focus on ideal customer, value prop resonance.
    2. **Sales Proposition & Differentiation:** Positioning, competitive comparison, cost transparency.
    3. **Conversion Path & CTAs:** CTA analysis, social proof, commitment to early stage customers.
    4. **Target Audience Relevance:** Language, emotional connection, narrative, success stories.
    5. **SEO & AI Visibility:** Technical SEO (H1/H2/Alt tags) AND AI Readiness (is content structured for AI answers?).
    6. **Strategic Opportunities:** High-level growth gaps.

    Group all software recommendations into a final "Tech Stack" section.
    
    Return ONLY a valid raw JSON object with this exact structure:
    {{
        "executive_summary": "<High-level summary of the store's potential.>",
        "score_breakdown": {{
             "score": <integer 1-10>,
             "reason": "<One sentence explaining the score.>"
        }},
        "section_1_branding": {{
            "title": "1. Branding & Messaging",
            "content": "<Analysis of clarity, consistency, and value prop resonance.>",
            "improvements": ["<Point 1>", "<Point 2>"]
        }},
        "section_2_sales": {{
            "title": "2. Sales Proposition & Differentiation",
            "content": "<Analysis of positioning, demonstration, and transparency.>",
            "improvements": ["<Point 1>", "<Point 2>"]
        }},
        "section_3_conversion": {{
            "title": "3. Conversion Path & CTAs",
            "content": "<Analysis of CTAs, social proof, and friction points.>",
            "improvements": ["<Point 1>", "<Point 2>"]
        }},
        "section_4_audience": {{
            "title": "4. Target Audience Relevance",
            "content": "<Analysis of emotional connection, language fit, and narrative.>",
            "improvements": ["<Point 1>", "<Point 2>"]
        }},
        "section_5_seo": {{
            "title": "5. SEO & AI Visibility",
            "content": "<Analysis of findability.>",
            "technical_notes": "H1: {data['h1_tags']} | Images: {data['image_stats']}",
            "ai_readiness": "<Is the content structured for AI/LLMs? (e.g. clear Q&A, lists)>",
            "improvements": ["<Point 1>", "<Point 2>"]
        }},
        "section_6_strategy": {{
            "title": "6. Strategic Opportunities",
            "content": "<Major growth levers not currently being used.>",
            "improvements": ["<Point 1>", "<Point 2>"]
        }},
        "recommended_stack": [
            {{
                "category": "SEO & Tech",
                "tool": "Plug In SEO",
                "service": "Technical SEO Audit & Fixes"
            }},
            {{
                "category": "Email Marketing",
                "tool": "Klaviyo",
                "service": "Flow Automation Setup"
            }},
            {{
                "category": "Social Proof",
                "tool": "Judge.me",
                "service": "Review Widget Customization"
            }}
        ]
    }}
    """
    try:
        response = model.generate_content(prompt)
        clean_json = response.text.strip().replace('```json', '').replace('```', '')
        audit_data = json.loads(clean_json)
        
        # --- NEW LOGIC: INJECT AFFILIATE LINKS ---
        # This loops through the 'recommended_stack' and adds a 'link' field
        if "recommended_stack" in audit_data:
            for item in audit_data['recommended_stack']:
                app_name = item.get('tool', '')
                # Default to Shopify App Store search if you don't have a specific link
                default_link = f"https://apps.shopify.com/search?q={app_name.replace(' ', '%20')}"
                
                # Check our dictionary for a match
                # We use a partial match check (e.g. if AI says "Klaviyo Email", it matches "Klaviyo")
                item['link'] = default_link # Start with default
                for key, affiliate_url in APP_LINKS.items():
                    if key.lower() in app_name.lower():
                        item['link'] = affiliate_url
                        break
        
        return audit_data
    except Exception as e:
        return {"error": str(e)}


def create_word_doc(audit, url):
    """Generates a formatted .docx file with the 6-point structure."""
    doc = Document()
    
    # Title
    heading = doc.add_heading(f'{AGENCY_NAME} Strategic Audit', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Audited Site: {url}")
    current_date = datetime.date.today().strftime("%B %d, %Y")
    doc.add_paragraph(f"Date: {current_date}")
    doc.add_paragraph("_" * 50) 
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(audit['executive_summary'])
    
    p_score = doc.add_paragraph()
    run_label = p_score.add_run("Marketing Health Score: ")
    run_label.bold = True
    run_score = p_score.add_run(f"{audit['score_breakdown']['score']}/10")
    run_score.bold = True
    run_score.font.color.rgb = RGBColor(0, 128, 96) # Green
    
    # --- LOOP THROUGH THE 6 SECTIONS ---
    sections = [
        audit['section_1_branding'],
        audit['section_2_sales'],
        audit['section_3_conversion'],
        audit['section_4_audience'],
        audit['section_5_seo'],
        audit['section_6_strategy']
    ]
    
    for sec in sections:
        doc.add_heading(sec['title'], level=1)
        doc.add_paragraph(sec['content'])
        
        # Add special sub-fields for SEO section
        if "technical_notes" in sec:
            doc.add_heading("Technical Data", level=2)
            doc.add_paragraph(sec['technical_notes'])
            doc.add_paragraph(f"AI Visibility: {sec['ai_readiness']}")

        doc.add_heading("Recommended Improvements:", level=3)
        for imp in sec['improvements']:
            doc.add_paragraph(imp, style='List Bullet')

    # 4. Recommended Stack
    doc.add_heading('7. Recommended Tools & Services', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Recommended Tool'
    hdr_cells[2].text = 'Inkroast Service'
    
    for stack in audit['recommended_stack']:
        row_cells = table.add_row().cells
        row_cells[0].text = stack['category']
        row_cells[1].text = stack['tool']
        row_cells[2].text = stack['service']

    # --- AGENCY PITCH SECTION ---
    doc.add_page_break()
    doc.add_heading('Turn this Audit into Revenue', level=1)
    doc.add_paragraph("You have the roadmap, now you need the execution. Inkroast specializes in implementing these exact strategies for Shopify brands.")
    
    p_call = doc.add_paragraph()
    p_call.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_call = p_call.add_run("\nBook your free implementation review:\n")
    run_call.bold = True
    run_link = p_call.add_run(BOOKING_LINK)
    run_link.font.color.rgb = RGBColor(0, 0, 255)
    run_link.underline = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 5. THE UI LAYOUT ---

st.title(f"☕ {AGENCY_NAME} Shopify e-commerce Snapshot")
st.markdown("### Complimentary Brand & SEO Audit")
st.markdown("Enter your Shopify Store URL. Our Agent APP will analyze your brand using the **Inkroast 6-Point Framework**.")

url_input = st.text_input("Store URL", placeholder="yourshopifystore.com", label_visibility="collapsed")

if st.button("Generate My Report 🚀", type="primary"):
    if not url_input:
        st.warning("Please enter a URL.")
    else:
        with st.spinner("Scanning H-Tags, Brand Voice, and Conversion Paths..."):
            data, error = scrape_shopify_store(url_input)
            
            if error:
                st.error(error)
            else:
                audit = analyze_store_json(data)
                
                if "error" in audit:
                    st.error("Analysis failed. Please try again.")
                else:
                    # --- PREVIEW ---
                    st.success("✅ Analysis Complete.")
                    
                    with st.container(border=True):
                        st.subheader("Executive Summary")
                        st.write(audit['executive_summary'])
                        st.metric("Health Score", f"{audit['score_breakdown']['score']}/10")
                        
                        # --- 4 TABS TO GROUP THE 6 POINTS ---
                        tab1, tab2, tab3, tab4 = st.tabs(["🎨 Brand & Sales", "👥 Audience & CTAs", "🔍 SEO & AI", "🛠 Strategy & Stack"])
                        
                        with tab1:
                            # Point 1
                            st.markdown(f"### {audit['section_1_branding']['title']}")
                            st.write(audit['section_1_branding']['content'])
                            for imp in audit['section_1_branding']['improvements']:
                                st.markdown(f"- {imp}")
                            st.markdown("---")
                            # Point 2
                            st.markdown(f"### {audit['section_2_sales']['title']}")
                            st.write(audit['section_2_sales']['content'])
                            for imp in audit['section_2_sales']['improvements']:
                                st.markdown(f"- {imp}")

                        with tab2:
                            # Point 3
                            st.markdown(f"### {audit['section_3_conversion']['title']}")
                            st.write(audit['section_3_conversion']['content'])
                            for imp in audit['section_3_conversion']['improvements']:
                                st.markdown(f"- {imp}")
                            st.markdown("---")
                            # Point 4
                            st.markdown(f"### {audit['section_4_audience']['title']}")
                            st.write(audit['section_4_audience']['content'])
                            for imp in audit['section_4_audience']['improvements']:
                                st.markdown(f"- {imp}")

                        with tab3:
                            # Point 5 (Detailed)
                            st.markdown(f"### {audit['section_5_seo']['title']}")
                            st.write(audit['section_5_seo']['content'])
                            
                            st.markdown("#### 🤖 AI Visibility Check")
                            st.info(audit['section_5_seo']['ai_readiness'])

                            st.markdown("#### ⚙️ Technical Data")
                            st.code(audit['section_5_seo']['technical_notes'], language="text")
                            
                            for imp in audit['section_5_seo']['improvements']:
                                st.markdown(f"- {imp}")
                                    
                        with tab4:
                            # Point 6
                            st.markdown(f"### {audit['section_6_strategy']['title']}")
                            st.write(audit['section_6_strategy']['content'])
                            st.markdown("---")
                            
                            st.markdown("### Recommended Tech Stack")
                            for stack in audit['recommended_stack']:
                                st.markdown(f"""
                                <div class="service-box">
                                    <div style="font-size:12px; text-transform:uppercase; color:#666; font-weight:bold; margin-bottom:5px;">{stack['category']}</div>
                                    <div style="display:flex; justify-content:space-between; align-items:center;">
                                        <div>
                                            <strong style="color:#333;">🛠 Tool:</strong> {stack['tool']}
                                        </div>
                                        <div>
                                            <strong style="color:#008060;">🚀 Service:</strong> {stack['service']}
                                        </div>
                                    </div>
                                </div>
                                """, unsafe_allow_html=True)
                            
                            st.markdown("### Ready to implement?")
                            st.link_button(
                                label="📅 Book a Discovery Call with Inkroast",
                                url=BOOKING_LINK,
                                use_container_width=True
                            )

                    # --- DOWNLOAD ---
                    doc_file = create_word_doc(audit, url_input)
                    
                    st.markdown("### 📥 Take this report with you")
                    st.download_button(
                        label="Download PDF Report (.docx)",
                        data=doc_file,
                        file_name=f"{AGENCY_NAME}_Audit_{url_input}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
